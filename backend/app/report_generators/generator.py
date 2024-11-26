import matplotlib.pyplot as plt
import seaborn as sns
import io
import base64
import json
import pandas as pd
import numpy as np
from typing import Dict, Any, List, Optional, Union
import pdfkit
from jinja2 import Template, Environment, FileSystemLoader
import plotly.express as px
import plotly.graph_objects as go
import plotly.io as pio
from datetime import datetime
import logging
import os
from pathlib import Path
from docx import Document
from ..core.exceptions import ReportGenerationError

# Configure logger
logger = logging.getLogger(__name__)

class ReportGenerator:
    def __init__(self, cohere_client=None):
        self.co = cohere_client
        self.logger = logger
        
        # Setup template environment
        template_dir = Path(__file__).parent.parent / 'templates'
        self.env = Environment(loader=FileSystemLoader(str(template_dir)))
        
        # Configure wkhtmltopdf path based on OS
        if os.name == 'nt':  # Windows
            self.wkhtmltopdf_path = r'C:\Program Files\wkhtmltopdf\bin\wkhtmltopdf.exe'
        else:  # Linux/Mac
            self.wkhtmltopdf_path = '/usr/local/bin/wkhtmltopdf'
        
        self.pdf_options = {
            'page-size': 'A4',
            'margin-top': '0.75in',
            'margin-right': '0.75in',
            'margin-bottom': '0.75in',
            'margin-left': '0.75in',
            'encoding': "UTF-8",
            'no-outline': None,
            'enable-local-file-access': None
        }

        # Configure plotly templates
        self.viz_templates = {
            'time_series': self._time_series_viz,
            'numerical': self._numerical_viz,
            'categorical': self._categorical_viz,
            'mixed': self._mixed_viz
        }

    def _create_plotly_figure(self, data: Dict[str, Any], viz_type: str) -> Dict[str, Any]:
        """Create a Plotly figure based on data type"""
        try:
            df = pd.DataFrame(data)
            
            if viz_type == 'time_series':
                fig = px.line(df, x=df.index, y=df.columns)
            elif viz_type == 'numerical':
                fig = px.scatter(df)
            elif viz_type == 'categorical':
                fig = px.bar(df)
            elif viz_type == 'mixed':
                fig = go.Figure()
                for col in df.columns:
                    if pd.api.types.is_numeric_dtype(df[col]):
                        fig.add_trace(go.Scatter(x=df.index, y=df[col], name=col))
                    else:
                        fig.add_trace(go.Bar(x=df.index, y=df[col], name=col))
            else:
                raise ValueError(f"Unsupported visualization type: {viz_type}")

            # Update layout
            fig.update_layout(
                template='plotly_white',
                title=f'{viz_type.title()} Visualization',
                showlegend=True,
                height=600,
                margin=dict(l=50, r=50, t=50, b=50)
            )

            return {
                'data': fig.to_dict()['data'],
                'layout': fig.to_dict()['layout']
            }

        except Exception as e:
            self.logger.error(f"Error creating Plotly figure: {str(e)}")
            raise ValueError(f"Failed to create visualization: {str(e)}")

    async def generate_report(self, data: Dict[str, Any], format: str = "json") -> Union[Dict[str, Any], bytes]:
        """Generate a report in the specified format"""
        try:
            if format == "json":
                return {
                    "status": "success",
                    "data": data,
                    "metadata": {
                        "generated_at": datetime.now().isoformat(),
                        "format": format
                    }
                }
            elif format == "pdf":
                template = self.env.get_template('report_template.html')
                html_content = template.render(
                    report=data,
                    dark_mode=True,
                    timestamp=datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                )
                
                config = pdfkit.configuration(wkhtmltopdf=self.wkhtmltopdf_path)
                pdf_content = pdfkit.from_string(
                    html_content,
                    False,
                    options=self.pdf_options,
                    configuration=config
                )
                return pdf_content
            elif format == "html":
                template = self.env.get_template('report_template.html')
                return template.render(
                    report=data,
                    dark_mode=True,
                    timestamp=datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                ).encode('utf-8')
            else:
                raise ValueError(f"Unsupported format: {format}")
                
        except Exception as e:
            self.logger.error(f"Report generation failed: {str(e)}")
            raise ReportGenerationError(f"Failed to generate {format} report: {str(e)}")

    def _generate_pdf_report(self, report_content: Dict[str, Any]) -> bytes:
        """Generate PDF report using report content"""
        try:
            # Load and render template
            template = self.env.get_template('report_template.html')
            html_content = template.render(
                report=report_content,
                dark_mode=True,
                timestamp=datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            )

            # Configure pdfkit with wkhtmltopdf path
            config = pdfkit.configuration(wkhtmltopdf=self.wkhtmltopdf_path)
            
            # Generate PDF
            pdf_content = pdfkit.from_string(
                html_content,
                False,  # Don't save to file, return bytes
                options=self.pdf_options,
                configuration=config
            )
            return pdf_content
        except Exception as e:
            raise ValueError(f"PDF generation failed: {str(e)}")

    def _generate_html_report(self, report_content: Dict[str, Any]) -> str:
        """Generate HTML report using report content"""
        try:
            return self._create_report_template(report_content)
        except Exception as e:
            raise ValueError(f"HTML generation failed: {str(e)}")

    def _create_report_template(self, report_content: Dict[str, Any]) -> str:
        """Create HTML template for report"""
        try:
            template = """
            <!DOCTYPE html>
            <html>
            <head>
                <title>Analysis Report</title>
                <style>
                    body { font-family: Arial, sans-serif; }
                    .container { max-width: 1200px; margin: 0 auto; padding: 20px; }
                    .section { margin-bottom: 30px; }
                    .visualization { max-width: 100%; height: auto; }
                </style>
            </head>
            <body>
                <div class="container">
                    <h1>Analysis Report</h1>
                    
                    <div class="section">
                        <h2>Analysis Results</h2>
                        <div>{{ analysis_results }}</div>
                    </div>
                    
                    {% if visualizations %}
                    <div class="section">
                        <h2>Visualizations</h2>
                        {% for viz in visualizations %}
                        <div class="visualization">{{ viz }}</div>
                        {% endfor %}
                    </div>
                    {% endif %}
                    
                    <div class="section">
                        <h2>Statistical Insights</h2>
                        <div>{{ statistical_insights }}</div>
                    </div>
                </div>
            </body>
            </html>
            """
            
            # Convert analysis results to HTML
            analysis_results = self._convert_to_html(report_content["analysis"])
            statistical_insights = self._convert_to_html(report_content.get("insights", {}))
            
            # Replace placeholders
            return template.replace(
                "{{ analysis_results }}", analysis_results
            ).replace(
                "{{ statistical_insights }}", statistical_insights
            ).replace(
                "{{ visualizations }}", self._format_visualizations(report_content.get("visualizations", []))
            )
            
        except Exception as e:
            raise ValueError(f"Template creation failed: {str(e)}")

    def _convert_to_html(self, data: Dict[str, Any]) -> str:
        """Convert dictionary data to HTML format"""
        html = "<dl>"
        for key, value in data.items():
            html += f"<dt><strong>{key}</strong></dt>"
            if isinstance(value, dict):
                html += f"<dd>{self._convert_to_html(value)}</dd>"
            else:
                html += f"<dd>{value}</dd>"
        html += "</dl>"
        return html

    def _format_visualizations(self, visualizations: List[str]) -> str:
        """Format visualizations for HTML display"""
        return "\n".join([f'<div class="viz">{viz}</div>' for viz in visualizations])

    def _generate_text_report(self, analysis: str) -> str:
        """Generate a simple report for text data"""
        template = Template("""
        <div class="report">
            <h1>Text Analysis Report</h1>
            <div class="analysis">
                <h2>Analysis</h2>
                {{ analysis | replace('\n', '<br>') }}
            </div>
        </div>
        """)
        return template.render(analysis=analysis)

    async def _generate_analysis(self, data: Dict[str, Any], query: str) -> str:
        """Generate analysis using Cohere"""
        try:
            prompt = self._generate_analysis_prompt(data, query)
            self.logger.info(f"Generated prompt: {prompt}")
            
            response = self.co.generate(
                prompt=prompt,
                max_tokens=1000,
                temperature=0.7,
                k=0,
                stop_sequences=[],
                return_likelihoods='NONE'
            )
            
            if not response.generations:
                raise Exception("No response generated from Cohere")
                
            return response.generations[0].text
        except Exception as e:
            self.logger.error(f"Error in _generate_analysis: {str(e)}")
            raise Exception(f"Analysis generation failed: {str(e)}")

    def _generate_analysis_prompt(self, data: Dict[str, Any], query: str) -> str:
        """Generate analysis prompt for Cohere"""
        try:
            if data.get('type') == 'unstructured':
                # Handle unstructured (text) data
                text_content = data.get('data', {}).get('text', '')
                return f"""Analyze the following text and answer this question: {query}

                Text Content:
                {text_content[:1000]}  # Limit text length

                Please provide:
                1. Key Points
                2. Main Themes
                3. Insights
                4. Recommendations

                Format the response in clear sections with bullet points."""
            else:
                # Handle structured data
                metadata = data.get('metadata', {})
                stats = metadata.get('statistics', {})
                
                return f"""Analyze the following data and answer this question: {query}

                Data Statistics:
                - Rows: {stats.get('row_count', 'N/A')}
                - Columns: {stats.get('column_count', 'N/A')}
                
                Data Sample:
                {json.dumps(data.get('data', [])[:5], indent=2)}
                
                Please provide:
                1. Key Insights
                2. Detailed Analysis
                3. Business Implications
                4. Limitations and Considerations
                
                Format the response in clear sections with bullet points."""
        except Exception as e:
            self.logger.error(f"Error generating prompt: {str(e)}")
            raise Exception(f"Failed to generate analysis prompt: {str(e)}")

    def _create_visualization(self, data: Dict[str, Any], query: str) -> Dict[str, str]:
        """Create appropriate visualizations based on data type"""
        try:
            df = pd.DataFrame(data['data'])
            data_type = data.get('data_type', 'mixed')
            metadata = data.get('metadata', {})
            
            viz_func = self.viz_templates.get(data_type, self._mixed_viz)
            return viz_func(df, metadata, query)
        except Exception as e:
            raise Exception(f"Error creating visualization: {str(e)}")

    def _time_series_viz(self, df: pd.DataFrame, metadata: Dict[str, Any], query: str) -> Dict[str, str]:
        """Generate time series visualizations"""
        date_cols = metadata.get('date_columns', [])
        if not date_cols:
            return self._mixed_viz(df, metadata, query)
        
        date_col = date_cols[0]
        numeric_cols = metadata.get('numeric_columns', [])
        
        # Interactive plot
        fig = px.line(df, x=date_col, y=numeric_cols[0] if numeric_cols else None)
        interactive_plot = pio.to_html(fig, full_html=False)
        
        # Static plot
        plt.figure(figsize=(10, 6))
        sns.lineplot(data=df, x=date_col, y=numeric_cols[0] if numeric_cols else None)
        static_plot = self._fig_to_base64()
        
        return {
            "interactive": interactive_plot,
            "static": static_plot
        }

    def _numerical_viz(self, df: pd.DataFrame, metadata: Dict[str, Any], query: str) -> Dict[str, str]:
        """Generate numerical visualizations"""
        numeric_cols = metadata.get('numeric_columns', [])
        if not numeric_cols:
            return self._mixed_viz(df, metadata, query)
        
        # Interactive plot
        fig = px.histogram(df, x=numeric_cols[0])
        interactive_plot = pio.to_html(fig, full_html=False)
        
        # Static plot
        plt.figure(figsize=(10, 6))
        sns.histplot(data=df, x=numeric_cols[0])
        static_plot = self._fig_to_base64()
        
        return {
            "interactive": interactive_plot,
            "static": static_plot
        }

    def _categorical_viz(self, df: pd.DataFrame, metadata: Dict[str, Any], query: str) -> Dict[str, str]:
        """Generate categorical visualizations"""
        cat_cols = metadata.get('categorical_columns', [])
        if not cat_cols:
            return self._mixed_viz(df, metadata, query)
        
        # Interactive plot
        fig = px.bar(df[cat_cols[0]].value_counts())
        interactive_plot = pio.to_html(fig, full_html=False)
        
        # Static plot
        plt.figure(figsize=(10, 6))
        sns.countplot(data=df, x=cat_cols[0])
        static_plot = self._fig_to_base64()
        
        return {
            "interactive": interactive_plot,
            "static": static_plot
        }

    def _mixed_viz(self, df: pd.DataFrame, metadata: Dict[str, Any], query: str) -> Dict[str, str]:
        """Generate mixed type visualizations"""
        # Interactive plot
        fig = go.Figure()
        for col in df.select_dtypes(include=[np.number]).columns[:3]:
            fig.add_trace(go.Box(y=df[col], name=col))
        interactive_plot = pio.to_html(fig, full_html=False)
        
        # Static plot
        plt.figure(figsize=(10, 6))
        df.select_dtypes(include=[np.number]).boxplot()
        static_plot = self._fig_to_base64()
        
        return {
            "interactive": interactive_plot,
            "static": static_plot
        }

    def _fig_to_base64(self) -> str:
        """Convert matplotlib figure to base64 string"""
        buffer = io.BytesIO()
        plt.savefig(buffer, format='png', bbox_inches='tight')
        plt.close()
        buffer.seek(0)
        return base64.b64encode(buffer.getvalue()).decode()

    def _generate_docx_report(self, analysis: str, visualizations: Dict[str, str], data: Dict[str, Any]) -> str:
        """Generate DOCX report"""
        doc = Document()
        doc.add_heading('Data Analysis Report', level=1)
        doc.add_heading('Analysis', level=2)
        doc.add_paragraph(analysis)

        # Add visualizations and metadata
        # ...

        # Save to a temporary file and return the path
        temp_path = '/tmp/report.docx'
        doc.save(temp_path)
        return temp_path